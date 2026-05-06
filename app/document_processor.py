from __future__ import annotations

import shutil
import tempfile
from copy import deepcopy
from decimal import Decimal
from pathlib import Path

import pdfplumber
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from instruction_parser import parse_instructions
from models import InstructionPlan, PriceMatch, ProcessOptions, ProcessResult
from money import format_money, format_percent, round_money
from price_finder import find_prices_in_text


REVISED_FOLDER_NAME = "Revised"
APP_TITLE = "Statement Markup Tool (FOR DAD)"
OUTPUT_REVIEW_NOTE = (
    "Review this output before sending or billing. The original file was not changed."
)
TOTAL_EXCLUDED_WORDS = ("total", "subtotal", "hst", "tax", "deposit", "balance")


def process_file(path: Path, instructions: str, options: ProcessOptions) -> ProcessResult:
    if not path.exists():
        raise ValueError("Select a document before processing.")

    plan = parse_instructions(instructions)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return write_revised_docx_from_pdf(path, plan, options)
    if suffix == ".docx":
        return write_revised_docx_from_docx(path, path, plan, options)
    if suffix == ".doc":
        with tempfile.TemporaryDirectory() as temp_dir:
            converted = convert_doc_to_docx(path, Path(temp_dir))
            return write_revised_docx_from_docx(path, converted, plan, options)
    raise ValueError("Please select a PDF or Word file.")


def convert_doc_to_docx(path: Path, output_dir: Path) -> Path:
    try:
        import pythoncom
        import win32com.client
    except ImportError as exc:
        raise ValueError("Legacy .doc files need Microsoft Word installed. Save the file as .docx and try again.") from exc

    converted_path = output_dir / f"{path.stem}.docx"
    pythoncom.CoInitialize()
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(str(path.resolve()), ReadOnly=True)
        document.SaveAs2(str(converted_path), FileFormat=16)
    except Exception as exc:
        raise ValueError("Could not read the .doc file. If Microsoft Word is not installed, save it as .docx and try again.") from exc
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()
    return converted_path


def get_revised_dir(input_path: Path) -> Path:
    revised_dir = input_path.parent / REVISED_FOLDER_NAME
    revised_dir.mkdir(exist_ok=True)
    return revised_dir


def build_output_path(input_path: Path, label: str, extension: str = ".docx") -> Path:
    revised_dir = get_revised_dir(input_path)
    candidate = revised_dir / f"{input_path.stem}_{label}{extension}"
    counter = 2
    while candidate.exists():
        candidate = revised_dir / f"{input_path.stem}_{label}_{counter}{extension}"
        counter += 1
    return candidate


def replace_prices_in_text(
    text: str,
    plan: InstructionPlan,
    options: ProcessOptions,
    context_hint: str = "",
) -> tuple[str, list[PriceMatch]]:
    matches = find_prices_in_text(text, plan, context_hint)
    if not matches:
        return text, []

    revised_parts: list[str] = []
    cursor = 0
    for match in matches:
        revised_parts.append(text[cursor : match.start])
        revised_parts.append(format_money(match.marked_up_value, options.round_to_whole_dollar))
        cursor = match.end
    revised_parts.append(text[cursor:])
    return "".join(revised_parts), matches


def replace_paragraph_prices(paragraph, plan: InstructionPlan, options: ProcessOptions, context_hint: str = "") -> list[PriceMatch]:
    revised_text, matches = replace_prices_in_text(paragraph.text, plan, options, context_hint)
    if matches:
        paragraph.text = revised_text
    return matches


def write_revised_docx_from_docx(input_path: Path, source_path: Path, plan: InstructionPlan, options: ProcessOptions) -> ProcessResult:
    output_path = build_output_path(input_path, "processed_statement")
    shutil.copy2(source_path, output_path)
    document = Document(str(output_path))
    all_matches: list[PriceMatch] = []

    for paragraph in document.paragraphs:
        all_matches.extend(replace_paragraph_prices(paragraph, plan, options))

    for table in document.tables:
        for row in table.rows:
            row_context = " ".join(cell.text for cell in row.cells)
            row_context_lines = build_row_context_lines(row)
            for cell in row.cells:
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    context_hint = row_context
                    if len(cell.paragraphs) > 1 and paragraph_index < len(row_context_lines):
                        context_hint = row_context_lines[paragraph_index]
                    all_matches.extend(replace_paragraph_prices(paragraph, plan, options, context_hint))

    net_total = calculate_net_total(all_matches, options)
    warnings = build_warnings(plan, all_matches)
    if all_matches and should_add_totals(plan):
        append_totals(document, plan, options, net_total)
    append_processing_summary(document, plan, options, all_matches, warnings)

    document.save(output_path)
    return ProcessResult(output_path=output_path, price_count=len(all_matches), net_total=net_total, warnings=tuple(warnings))


def write_revised_docx_from_pdf(input_path: Path, plan: InstructionPlan, options: ProcessOptions) -> ProcessResult:
    output_path = build_output_path(input_path, "processed_statement")
    document = Document()
    configure_document_styles(document)
    document.add_heading("Processed Statement", level=1)
    document.add_paragraph(f"Source file: {input_path.name}")
    document.add_paragraph(OUTPUT_REVIEW_NOTE)

    all_matches: list[PriceMatch] = []
    with pdfplumber.open(input_path) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            if page_index > 1:
                document.add_page_break()
            document.add_heading(f"Page {page_index}", level=2)
            text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
            revised_text, matches = replace_prices_in_text(text, plan, options)
            all_matches.extend(matches)
            add_pdf_text(document, revised_text)

    net_total = calculate_net_total(all_matches, options)
    warnings = build_warnings(plan, all_matches)
    if all_matches and should_add_totals(plan):
        append_totals(document, plan, options, net_total)
    append_processing_summary(document, plan, options, all_matches, warnings)

    document.save(output_path)
    return ProcessResult(output_path=output_path, price_count=len(all_matches), net_total=net_total, warnings=tuple(warnings))


def build_row_context_lines(row) -> list[str]:
    lines: list[str] = []
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for raw_line in paragraph.text.splitlines():
                text = " ".join(raw_line.split())
                if text and any(character.isalpha() for character in text):
                    lines.append(text)
    return lines


def configure_document_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = document.styles
    styles["Normal"].font.name = "Segoe UI"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Segoe UI"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Segoe UI"
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.bold = True


def add_pdf_text(document: Document, text: str) -> None:
    for line in text.splitlines() or [""]:
        cleaned = line.strip()
        if not cleaned:
            document.add_paragraph()
            continue
        paragraph = document.add_paragraph(cleaned)
        paragraph.paragraph_format.space_after = Pt(2)


def calculate_net_total(matches: list[PriceMatch], options: ProcessOptions) -> Decimal:
    return sum(
        (round_money(match.marked_up_value, options.round_to_whole_dollar) for match in matches if include_in_net_total(match)),
        Decimal("0"),
    )


def include_in_net_total(match: PriceMatch) -> bool:
    combined = f"{match.label} {match.context}".lower()
    return not any(word in combined for word in TOTAL_EXCLUDED_WORDS)


def should_add_totals(plan: InstructionPlan) -> bool:
    totals = plan.totals
    return any((totals.add_net_total, totals.hst_rate is not None, totals.add_grand_total, totals.add_deposit, totals.add_balance))


def append_totals(document: Document, plan: InstructionPlan, options: ProcessOptions, net_total: Decimal) -> None:
    rows = build_total_rows(plan, options, net_total)
    if not rows:
        return

    if document.tables:
        append_rows_to_existing_table(document.tables[-1], rows)
        return

    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.style = "Table Grid"
    for label, value, is_strong in rows:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].width = Inches(2.6)
        row.cells[1].width = Inches(1.8)
        style_total_cells(row.cells, is_strong)


def build_total_rows(plan: InstructionPlan, options: ProcessOptions, net_total: Decimal) -> list[tuple[str, str, bool]]:
    totals = plan.totals
    rows: list[tuple[str, str, bool]] = []
    running_total = net_total

    if totals.add_net_total:
        rows.append(("Net Total", format_money(net_total, options.round_to_whole_dollar), True))

    if totals.hst_rate is not None:
        hst_amount = round_money(net_total * totals.hst_rate, options.round_to_whole_dollar)
        rows.append((f"HST ({format_percent(totals.hst_rate)})", format_money(hst_amount, options.round_to_whole_dollar), False))
        running_total = round_money(net_total + hst_amount, options.round_to_whole_dollar)

    if totals.add_grand_total:
        rows.append(("Total", format_money(running_total, options.round_to_whole_dollar), True))

    if totals.add_deposit:
        deposit = calculate_deposit(totals.deposit_rate, totals.deposit_amount, totals.deposit_adjustment, running_total, options)
        rows.append(("Deposit", format_money(deposit, options.round_to_whole_dollar), False))
        if totals.add_balance:
            balance = round_money(running_total - deposit, options.round_to_whole_dollar)
            rows.append(("Balance", format_money(balance, options.round_to_whole_dollar), True))
    elif totals.add_balance:
        rows.append(("Balance", format_money(running_total, options.round_to_whole_dollar), True))

    return rows


def calculate_deposit(
    rate: Decimal | None,
    amount: Decimal | None,
    adjustment: Decimal,
    total: Decimal,
    options: ProcessOptions,
) -> Decimal:
    if amount is not None:
        return round_money(amount + adjustment, options.round_to_whole_dollar)
    if rate is not None:
        return round_money((total * rate) + adjustment, options.round_to_whole_dollar)
    return round_money(adjustment, options.round_to_whole_dollar)


def append_rows_to_existing_table(table, rows: list[tuple[str, str, bool]]) -> None:
    for label, value, is_strong in rows:
        row = table.add_row()
        copy_previous_row_format(table, row)
        row.cells[0].text = label
        if len(row.cells) > 1:
            for cell in row.cells[1:-1]:
                cell.text = ""
            row.cells[-1].text = value
        style_total_cells(row.cells, is_strong)


def copy_previous_row_format(table, row) -> None:
    if len(table.rows) < 2:
        return
    previous_row = table.rows[-2]
    for source_cell, target_cell in zip(previous_row.cells, row.cells):
        source_properties = source_cell._tc.get_or_add_tcPr()
        target_properties = target_cell._tc.get_or_add_tcPr()
        target_properties.attrib.clear()
        for child in list(target_properties):
            target_properties.remove(child)
        for name, value in source_properties.attrib.items():
            target_properties.set(name, value)
        for child in source_properties:
            target_properties.append(deepcopy(child))
        if source_cell.paragraphs and target_cell.paragraphs:
            target_cell.paragraphs[0].style = source_cell.paragraphs[0].style


def style_total_cells(cells, strong: bool) -> None:
    for index, cell in enumerate(cells):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if strong:
            shade_cell(cell, "EAF2F8")
        for paragraph in cell.paragraphs:
            if index == len(cells) - 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = strong


def shade_cell(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def append_processing_summary(
    document: Document,
    plan: InstructionPlan,
    options: ProcessOptions,
    matches: list[PriceMatch],
    warnings: list[str],
) -> None:
    document.add_paragraph()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Processing Summary")
    run.bold = True
    run.font.color.rgb = RGBColor(80, 80, 80)

    summary = document.add_paragraph()
    apply_style_if_available(summary, "Normal")
    summary.add_run(f"Prices updated: {len(matches)}. ")
    summary.add_run(f"Default multiplier: {plan.default_multiplier}. ")
    summary.add_run("Rounding: nearest whole dollar." if options.round_to_whole_dollar else "Rounding: cents preserved.")
    for warning in warnings:
        add_warning_paragraph(document, warning)


def apply_style_if_available(paragraph, style_name: str) -> None:
    try:
        paragraph.style = style_name
    except KeyError:
        return


def add_warning_paragraph(document: Document, warning: str) -> None:
    paragraph = document.add_paragraph()
    apply_style_if_available(paragraph, "List Bullet")
    if paragraph.style.name == "Normal":
        paragraph.add_run("- ")
    paragraph.add_run(warning)


def build_warnings(plan: InstructionPlan, matches: list[PriceMatch]) -> list[str]:
    warnings: list[str] = []
    if plan.default_multiplier == Decimal("1"):
        warnings.append("No default multiplier was found, so unmatched prices were left at 1.00x.")
    for rule in plan.specific_rules:
        if not any(match.multiplier == rule.multiplier for match in matches):
            warnings.append(f"No detected price appeared to match the special rule for '{rule.label}'.")
    if plan.totals.add_deposit and plan.totals.deposit_rate == Decimal("0.50"):
        warnings.append("Deposit was requested without an amount or percent, so the app used 50%.")
    if not matches:
        warnings.append("No prices were detected. Review the source document layout and instructions.")
    return warnings
